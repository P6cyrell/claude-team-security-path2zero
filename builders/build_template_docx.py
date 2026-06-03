"""
Build _template-leadership-update.docx — Arrive-branded leadership update.

Output is a .docx that, when uploaded to Google Drive, auto-converts to a
Google Doc with brand colors, font, and `{{PLACEHOLDER}}` tokens intact.
The renderer later does drive.files.copy on the converted Doc and
replaceAllText to fill placeholders.

Run:  python3 build_template_docx.py
"""

from __future__ import annotations

import json
import pathlib

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


SKILL_DIR = pathlib.Path(__file__).resolve().parent.parent
BRAND_DIR = SKILL_DIR / "brand"
OUT_DIR   = SKILL_DIR / "templates"
OUT_DIR.mkdir(exist_ok=True)

PALETTE = json.loads((BRAND_DIR / "palette.json").read_text())
FONT    = json.loads((BRAND_DIR / "font.json").read_text())

FONT_NAME = FONT["primary_name"]


def hex2rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


C = {k: hex2rgb(v) for k, v in PALETTE.items() if isinstance(v, str) and v.startswith("#")}


def shade_cell(cell, hex_no_hash: str):
    """Set table cell background fill."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_no_hash)
    tc_pr.append(shd)


def styled_run(paragraph, text, *, size=11, bold=False, color=None, font=FONT_NAME):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return r


def heading(doc, text, level: int):
    p = doc.add_paragraph()
    sizes = {1: 20, 2: 14, 3: 12}
    styled_run(p, text, size=sizes.get(level, 11), bold=True, color=C["primary_purple"])
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, *, bold=False):
    p = doc.add_paragraph()
    styled_run(p, text, size=11, bold=bold, color=C["dark_purple"])
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    styled_run(p, text, size=11, color=C["dark_purple"])
    return p


def make_table(doc, headers: list[str], placeholder_rows: int, col_widths_in: list[float]):
    tbl = doc.add_table(rows=1 + placeholder_rows, cols=len(headers))
    tbl.autofit = False
    tbl.allow_autofit = False

    # Header
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.width = Inches(col_widths_in[j])
        shade_cell(cell, "5F016F")
        cell.text = ""
        p = cell.paragraphs[0]
        styled_run(p, h, size=11, bold=True, color=C["white"])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for i in range(1, placeholder_rows + 1):
        for j in range(len(headers)):
            cell = tbl.rows[i].cells[j]
            cell.width = Inches(col_widths_in[j])
            shade_cell(cell, "F9F5F4" if i % 2 == 0 else "FFFFFF")

    return tbl


def main():
    doc = Document()

    # Set base body font + page margins
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # ----- header band -----
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run()
    run.add_picture(str(BRAND_DIR / "arrive-wordmark.png"), height=Inches(0.45))

    # Title block
    p = doc.add_paragraph()
    styled_run(p, "{{TEAM_NAME}} — Path to Zero", size=22, bold=True,
               color=C["primary_purple"])
    p = doc.add_paragraph()
    styled_run(p, "Critical Vulnerability Remediation  ·  Security Leadership Update  ·  {{AS_OF_DATE}}",
               size=12, color=C["muted_purple"])
    p = doc.add_paragraph()
    styled_run(p, "Engineering lead: {{ENGINEERING_LEAD}}    Q2 epics: {{EPIC_LIST}}",
               size=11, color=C["dark_purple"])

    # Accent rule
    rule = doc.add_paragraph()
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "FFADE4")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ----- §1 Executive Summary -----
    heading(doc, "1. Executive Summary", 1)
    body(doc, "Headline figures as of {{AS_OF_DATE}}:")
    bullet(doc, "Peak open criticals: {{STAT_PEAK}}")
    bullet(doc, "Today: {{STAT_TODAY}}")
    bullet(doc, "Services with criticals: {{STAT_SERVICES_WITH_CRITS}}")
    bullet(doc, "Services after EOD: {{STAT_SERVICES_AFTER}}")
    body(doc, "Narrative:")
    bullet(doc, "{{HEADLINE_BULLETS}}")

    # ----- §2 Burndown -----
    heading(doc, "2. Burndown", 1)
    body(doc, "The path-to-zero burndown shows three lines: actual progress (solid), "
              "original-plan baseline (dotted), and current forecast (dashed). "
              "The chart below is regenerated each run from the latest data.")
    # Bookmark for burndown image insertion
    p = doc.add_paragraph()
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), "0")
    bm_start.set(qn("w:name"), "burndown")
    p._p.append(bm_start)
    styled_run(p, "[ {{BURNDOWN_PLACEHOLDER}} ]", size=10, color=C["muted_purple"])
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), "0")
    p._p.append(bm_end)

    # ----- §3 Remaining surface -----
    heading(doc, "3. Remaining attack surface", 1)
    body(doc, "Open critical-bearing repos as of {{AS_OF_DATE}}:")
    tbl = make_table(doc,
                     headers=["Repo", "Critical", "High", "Status"],
                     placeholder_rows=5,
                     col_widths_in=[2.9, 0.8, 0.8, 2.2])
    for i in range(1, 6):
        cells = tbl.rows[i].cells
        for j, ph in enumerate([
            "{{ROW_%d_REPO}}" % i,
            "{{ROW_%d_CRITICAL}}" % i,
            "{{ROW_%d_HIGH}}" % i,
            "{{ROW_%d_STATUS}}" % i,
        ]):
            cells[j].text = ""
            styled_run(cells[j].paragraphs[0], ph, size=10, color=C["dark_purple"])

    # ----- §4 Dependency status -----
    heading(doc, "4. {{DEP_GRID_TITLE}}", 1)
    body(doc, "{{DEP_GRID_NARRATIVE}}")
    dep_tbl = make_table(doc,
                         headers=["Region", "Status"],
                         placeholder_rows=3,
                         col_widths_in=[2.0, 4.7])
    for i in range(1, 4):
        cells = dep_tbl.rows[i].cells
        for j, ph in enumerate([
            "{{REGION_%d_NAME}}" % i,
            "{{REGION_%d_STATUS}}" % i,
        ]):
            cells[j].text = ""
            styled_run(cells[j].paragraphs[0], ph, size=10, color=C["dark_purple"])

    # ----- §5 Critical path -----
    heading(doc, "5. Critical path — {{CP_SERVICE}}", 1)
    body(doc, "{{CP_STAT_LINE}}", bold=True)
    body(doc, "Why this is the critical path:")
    bullet(doc, "{{CP_RATIONALE_BULLETS}}")

    # ----- §6 Q2 epics -----
    heading(doc, "6. Q2 epic structure", 1)
    epic_tbl = make_table(doc,
                          headers=["Epic", "State", "Summary"],
                          placeholder_rows=3,
                          col_widths_in=[1.4, 1.4, 3.9])
    for i in range(1, 4):
        cells = epic_tbl.rows[i].cells
        for j, ph in enumerate([
            "{{EPIC_%d_ID}}" % i,
            "{{EPIC_%d_STATE}}" % i,
            "{{EPIC_%d_SUMMARY}}" % i,
        ]):
            cells[j].text = ""
            styled_run(cells[j].paragraphs[0], ph, size=10, color=C["dark_purple"])

    # ----- §7 Reallocation -----
    heading(doc, "7. Resource reallocation", 1)
    body(doc, "Engineers reallocated to security remediation (still in force):")
    realloc_tbl = make_table(doc,
                             headers=["Engineer", "From", "Allocation"],
                             placeholder_rows=5,
                             col_widths_in=[2.2, 2.3, 2.2])
    for i in range(1, 6):
        cells = realloc_tbl.rows[i].cells
        for j, ph in enumerate([
            "{{REALLOC_%d_ENGINEER}}" % i,
            "{{REALLOC_%d_FROM}}" % i,
            "{{REALLOC_%d_TO}}" % i,
        ]):
            cells[j].text = ""
            styled_run(cells[j].paragraphs[0], ph, size=10, color=C["dark_purple"])

    # ----- §8 Timeline -----
    heading(doc, "8. Updated timeline to zero", 1)
    tl_tbl = make_table(doc,
                        headers=["Date", "Milestone", "Open criticals"],
                        placeholder_rows=7,
                        col_widths_in=[1.2, 4.2, 1.3])
    for i in range(1, 8):
        cells = tl_tbl.rows[i].cells
        for j, ph in enumerate([
            "{{TIMELINE_%d_DATE}}" % i,
            "{{TIMELINE_%d_MILESTONE}}" % i,
            "{{TIMELINE_%d_OPEN}}" % i,
        ]):
            cells[j].text = ""
            styled_run(cells[j].paragraphs[0], ph, size=10, color=C["dark_purple"])

    # ----- §9 Asks -----
    heading(doc, "9. What we need from Security leadership", 1)
    for i in range(1, 5):
        p = doc.add_paragraph()
        styled_run(p, "{{ASK_%d_NUM}}. {{ASK_%d_TITLE}}" % (i, i),
                   size=12, bold=True, color=C["primary_purple"])
        body(doc, "{{ASK_%d_BODY}}" % i)

    # ----- Footer line -----
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(foot, "Confidential  ·  {{TEAM_NAME}}  ·  {{AS_OF_DATE}}",
               size=9, color=C["muted_purple"])

    out = OUT_DIR / "_template-leadership-update.docx"
    doc.save(str(out))
    print(f"wrote {out}  ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
